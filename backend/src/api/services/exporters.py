"""Format-specific exporters for generated legal documents."""

from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from typing import Protocol

from src.api.db.models import Case, GeneratedDocument
from src.api.schemas.exporting import ExportFormat


@dataclass(frozen=True)
class ExportArtifact:
    filename: str
    content_type: str
    content: bytes


@dataclass(frozen=True)
class ExportContext:
    document: GeneratedDocument
    case: Case
    include_citations: bool


class DocumentExporter(Protocol):
    format: ExportFormat

    def export(self, context: ExportContext) -> ExportArtifact: ...


def _metadata_lines(context: ExportContext) -> list[str]:
    document = context.document
    generated_at = getattr(document, "created_at", None) or datetime.now(UTC)
    return [
        f"Case: {context.case.title}",
        f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M UTC')}",
        f"Version: {document.version}",
    ]


def _citation_lines(context: ExportContext) -> list[str]:
    if not context.include_citations or not context.document.citations_data:
        return []
    return ["Citations:", *[f"- {source_id}" for source_id in context.document.citations_data]]


class MarkdownExporter:
    format = ExportFormat.MARKDOWN

    def export(self, context: ExportContext) -> ExportArtifact:
        sections = [f"# {context.document.title}", "", *(_metadata_lines(context)), "", context.document.content]
        citations = _citation_lines(context)
        if citations:
            sections.extend(["", *citations])
        return ExportArtifact(
            filename=f"{context.document.document_key}-v{context.document.version}.md",
            content_type="text/markdown",
            content="\n".join(sections).encode("utf-8"),
        )


class DocxExporter:
    format = ExportFormat.DOCX

    def export(self, context: ExportContext) -> ExportArtifact:
        from docx import Document
        from docx.shared import Inches, Pt

        document = Document()
        section = document.sections[0]
        section.top_margin = section.bottom_margin = Inches(1)
        title = document.add_paragraph()
        title_run = title.add_run(context.document.title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        for line in _metadata_lines(context):
            document.add_paragraph(line)
        document.add_paragraph()
        for paragraph in context.document.content.split("\n\n"):
            document.add_paragraph(paragraph)
        citations = _citation_lines(context)
        if citations:
            document.add_heading("Citations", level=2)
            for citation in citations[1:]:
                document.add_paragraph(citation[2:], style="List Bullet")
        output = BytesIO()
        document.save(output)
        return ExportArtifact(
            filename=f"{context.document.document_key}-v{context.document.version}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=output.getvalue(),
        )


class PdfExporter:
    format = ExportFormat.PDF

    def export(self, context: ExportContext) -> ExportArtifact:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        output = BytesIO()
        styles = getSampleStyleSheet()
        story = [Paragraph(context.document.title, styles["Title"]), Spacer(1, 12)]
        story.extend(Paragraph(line, styles["Normal"]) for line in _metadata_lines(context))
        story.append(Spacer(1, 12))
        story.extend(
            Paragraph(paragraph.replace("\n", "<br/>"), styles["BodyText"])
            for paragraph in context.document.content.split("\n\n")
        )
        citations = _citation_lines(context)
        if citations:
            story.extend([Spacer(1, 12), Paragraph("Citations", styles["Heading2"])])
            story.extend(Paragraph(line[2:], styles["Normal"]) for line in citations[1:])
        SimpleDocTemplate(output, pagesize=LETTER).build(story)
        return ExportArtifact(
            filename=f"{context.document.document_key}-v{context.document.version}.pdf",
            content_type="application/pdf",
            content=output.getvalue(),
        )


class ExporterRegistry:
    def __init__(self, exporters: list[DocumentExporter] | None = None) -> None:
        exporters = exporters or [MarkdownExporter(), DocxExporter(), PdfExporter()]
        self._exporters = {exporter.format: exporter for exporter in exporters}

    def get(self, export_format: ExportFormat) -> DocumentExporter:
        try:
            return self._exporters[export_format]
        except KeyError as exc:
            raise ValueError(f"Unsupported export format: {export_format.value}") from exc
